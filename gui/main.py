from sklearn import datasets
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.linear_model import SGDClassifier
from sklearn.pipeline import Pipeline
import numpy as np
import docx
from docx.shared import RGBColor
import pickle


def readFile(fileName):
    f = open(fileName, "r", encoding="utf-8")
    return f.readlines()


class Classifier:

    def __init__(self):
        with open('../svm_clf.pickle', 'rb') as f:
            self.svm_clf = pickle.load(f)

            self.predicted = []
            self.lines = []
            self.target_names = ['N', 'O', 'P']

    def getClass(self, index):
        return self.target_names[index]

    def predict(self, line):
        self.lines.append(line)
        self.predicted.append(self.svm_clf.predict([line])[0])

        return self.predicted[-1]

    def createDocx(self, filename):
        doc = docx.Document()

        for class_num, line in zip(self.predicted, self.lines):

            if(self.getClass(class_num) == 'N'):
                run = doc.add_paragraph().add_run()
                run.text = line
                run.font.underline = True
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x99, 0x33)
            else:
                parag = doc.add_paragraph(line)

        doc.save("".join(filename.split(".")[:-1]) + '.docx')


if __name__ == "__main__":
    lines = readFile("test.txt")

    cls = Classifier()
    for line in lines:
        cls.predict(line)

    cls.createDocx()
    print("Done")
